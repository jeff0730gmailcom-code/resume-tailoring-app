"""CV text extraction service.

Extracts raw text from an uploaded PDF or DOCX master CV so it can be sent
to the AI tailoring step.
"""
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from pypdf import PdfReader


class CvParsingError(Exception):
    """Raised when a CV file cannot be parsed into text."""


def extract_text_from_cv(file_path: Path, document: DocumentObject | None = None) -> str:
    """Extract plain text from a PDF or DOCX file.

    document, when the caller already loaded the DOCX (e.g. to also run
    style extraction / bullet-count detection on it — see
    app/api/routes/resume.py's upload handler), avoids re-parsing the same
    file from disk a second time.

    Raises CvParsingError if the file type is unsupported or extraction
    yields no usable text (e.g. a scanned/image-only PDF).
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_from_pdf(file_path)
    elif suffix == ".docx":
        text = _extract_from_docx(file_path, document)
    else:
        raise CvParsingError(f"Unsupported file type: {suffix}")

    text = text.strip()
    if not text:
        raise CvParsingError(
            "No text could be extracted from this file. If it's a scanned "
            "PDF (image-only), please upload a text-based PDF or DOCX instead."
        )
    return text


def _extract_from_pdf(file_path: Path) -> str:
    try:
        reader = PdfReader(str(file_path))
    except Exception as exc:  # noqa: BLE001 - surface as a parsing error
        raise CvParsingError(f"Could not read PDF file: {exc}") from exc

    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    return "\n".join(pages_text)


def _extract_from_docx(file_path: Path, document: DocumentObject | None = None) -> str:
    try:
        doc = document if document is not None else Document(str(file_path))
    except Exception as exc:  # noqa: BLE001 - surface as a parsing error
        raise CvParsingError(f"Could not read DOCX file: {exc}") from exc

    paragraphs = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs)
